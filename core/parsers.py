import io, re
from pathlib import Path
from docx import Document
from pypdf import PdfReader

def parse_text_file(filename, data):
    ext = Path(filename).suffix.lower()
    if ext == ".txt":
        text = data.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages)
    elif ext == ".docx":
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows: parts.append(" | ".join(c.text for c in row.cells))
        text = "\n".join(parts)
    else: raise ValueError("Unsupported file type")
    if not text.strip(): raise ValueError("No extractable text")
    return text.strip()

def extract_name(text, filename):
    for line in [x.strip() for x in text.splitlines() if x.strip()][:8]:
        x = re.sub(r"^(name|candidate name)\s*[:\-]\s*", "", line, flags=re.I)
        if 2 <= len(x.split()) <= 5 and re.fullmatch(r"[A-Za-z][A-Za-z .'-]+", x):
            if not any(k in x.lower() for k in ["resume","curriculum","email","phone"]): return x
    return Path(filename).stem.replace("_"," ").replace("-"," ")

def extract_contact(text):
    e = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    p = re.search(r"(?:\+?\d[\d ()-]{8,}\d)", text)
    return {"email": e.group(0) if e else "", "phone": p.group(0).strip() if p else ""}

def extract_experience(text):
    pats = [r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:professional\s+)?experience",
            r"experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)"]
    vals = [float(m.group(1)) for pat in pats for m in re.finditer(pat,text,re.I)]
    return max(vals) if vals else None

def extract_education(text):
    keys=("bachelor","b.e.","b.tech","btech","master","m.e.","m.tech","mtech","phd",
          "b.sc","bsc","m.sc","msc","degree","diploma","computer science",
          "artificial intelligence","data science")
    lines=[x.strip() for x in text.splitlines() if x.strip()]
    return " | ".join(x for x in lines if any(k in x.lower() for k in keys))[:1500]

def parse_resume_file(filename,data):
    text=parse_text_file(filename,data); c=extract_contact(text)
    return {"source_file":filename,"candidate_name":extract_name(text,filename),
            "email":c["email"],"phone":c["phone"],"experience_years":extract_experience(text),
            "education_summary":extract_education(text),"text":text}
